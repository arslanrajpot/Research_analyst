from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import and_, desc
from typing import List, Optional
import os
import shutil
import uuid
from datetime import datetime
import logging

from models.user import User, Document as DBDocument
from database import get_db
from auth import get_current_active_user
from config import UPLOAD_DIR

router = APIRouter(prefix="/documents", tags=["documents"])
logger = logging.getLogger(__name__)

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    filename: str = Form(...),
    file_type: str = Form(...),
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """
    Upload a document for analysis and storage.
    """
    try:
        # Validate file type
        allowed_extensions = {
            '.pdf', '.docx', '.doc', '.pptx', '.ppt', 
            '.xlsx', '.xls', '.csv', '.txt', '.md',
            '.jpg', '.jpeg', '.png', '.gif', '.bmp'
        }
        
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not supported. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Create database record
        document = DBDocument(
            user_id=current_user.id,
            filename=filename,
            file_path=file_path,
            file_type=file_type,
            file_size=os.path.getsize(file_path),
            status="uploaded",
            uploaded_at=datetime.utcnow()
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document uploaded: {filename} by user {current_user.id}")
        
        return {
            "document_id": str(document.id),
            "filename": filename,
            "file_type": file_type,
            "file_size": document.file_size,
            "status": "uploaded",
            "uploaded_at": document.uploaded_at.isoformat()
        }
        
    except Exception as e:
        logger.error(f"Document upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")

@router.get("/list")
async def list_documents(
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db),
    limit: int = 50,
    offset: int = 0
):
    """
    List user's uploaded documents.
    """
    try:
        documents = db.query(DBDocument).filter(
            and_(
                DBDocument.user_id == current_user.id,
                DBDocument.deleted_at.is_(None)
            )
        ).order_by(desc(DBDocument.uploaded_at)).offset(offset).limit(limit).all()
        
        return {
            "documents": [
                {
                    "id": str(doc.id),
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "file_size": doc.file_size,
                    "status": doc.status,
                    "uploaded_at": doc.uploaded_at.isoformat(),
                    "processed_at": doc.processed_at.isoformat() if doc.processed_at else None
                }
                for doc in documents
            ],
            "total": len(documents)
        }
        
    except Exception as e:
        logger.error(f"List documents error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")

@router.get("/{document_id}")
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """
    Get document details and content.
    """
    try:
        document = db.query(DBDocument).filter(
            and_(
                DBDocument.id == document_id,
                DBDocument.user_id == current_user.id,
                DBDocument.deleted_at.is_(None)
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "id": str(document.id),
            "filename": document.filename,
            "file_type": document.file_type,
            "file_size": document.file_size,
            "status": document.status,
            "uploaded_at": document.uploaded_at.isoformat(),
            "processed_at": document.processed_at.isoformat() if document.processed_at else None,
            "content": document.content if document.content else None,
            "metadata": document.file_metadata if document.file_metadata else {}
        }
        
    except Exception as e:
        logger.error(f"Get document error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")

@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """
    Delete a document (soft delete).
    """
    try:
        document = db.query(DBDocument).filter(
            and_(
                DBDocument.id == document_id,
                DBDocument.user_id == current_user.id,
                DBDocument.deleted_at.is_(None)
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Soft delete
        document.deleted_at = datetime.utcnow()
        db.commit()
        
        logger.info(f"Document deleted: {document.filename} by user {current_user.id}")
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        logger.error(f"Delete document error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@router.post("/{document_id}/process")
async def process_document(
    document_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """
    Process a document to extract text content and metadata.
    """
    try:
        document = db.query(DBDocument).filter(
            and_(
                DBDocument.id == document_id,
                DBDocument.user_id == current_user.id,
                DBDocument.deleted_at.is_(None)
            )
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        if document.status == "processed":
            return {"message": "Document already processed"}
        
        # Update status to processing
        document.status = "processing"
        db.commit()
        
        try:
            # Extract text content based on file type
            content = await extract_document_content(document.file_path, document.file_type)
            
            # Update document with extracted content
            document.content = content
            document.status = "processed"
            document.processed_at = datetime.utcnow()
            db.commit()
            
            logger.info(f"Document processed: {document.filename}")
            
            return {
                "message": "Document processed successfully",
                "content_length": len(content) if content else 0
            }
            
        except Exception as e:
            document.status = "error"
            db.commit()
            raise e
            
    except Exception as e:
        logger.error(f"Process document error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")

async def extract_document_content(file_path: str, file_type: str) -> str:
    """
    Extract text content from various document types.
    """
    try:
        if file_type.startswith('text/'):
            # Text files
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == 'application/pdf':
            # PDF files
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                logger.warning("PyPDF2 not installed, skipping PDF text extraction")
                return ""
        
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            # Word documents
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                logger.warning("python-docx not installed, skipping Word text extraction")
                return ""
        
        elif file_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                          'application/vnd.ms-excel']:
            # Excel files
            try:
                import pandas as pd
                df = pd.read_excel(file_path)
                return df.to_string()
            except ImportError:
                logger.warning("pandas not installed, skipping Excel text extraction")
                return ""
        
        else:
            logger.warning(f"Unsupported file type for text extraction: {file_type}")
            return ""
            
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return ""
